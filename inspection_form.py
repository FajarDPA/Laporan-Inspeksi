import io
import math
from datetime import date
import streamlit as st
from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import Table
import subprocess
import sys
import requests # New library to fetch files from URL

# Pustaka untuk mengirim email
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

# Pustaka untuk konversi PDF
# Catatan: Konversi PDF asli tidak didukung di lingkungan ini.
# Hanya unduhan DOCX yang tersedia.

# -----------------------------
# Helpers for working with docx
# -----------------------------

def _replace_in_paragraph(paragraph, placeholder, value):
    if placeholder not in paragraph.text:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text.replace(placeholder, value)
    for _ in range(len(paragraph.runs)):
        paragraph.runs[0].clear()
        paragraph.runs[0].text = ""
    if len(paragraph.runs) == 0:
        paragraph.add_run(new_text)
    else:
        paragraph.runs[0].text = new_text

def replace_placeholder_everywhere(doc: Document, placeholder: str, value: str):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, placeholder, value)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, placeholder, value)

def find_cell_with_text(doc: Document, placeholder: str):
    for tbl in doc.tables:
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        return cell, tbl, r_idx, c_idx
    return None, None, None, None

def _get_table_grid_col_widths_in_inches(tbl: Table):
    try:
        grid = tbl._tbl.tblGrid
        if grid is None:
            return None
        cols = []
        for gcol in grid.gridCol_lst:
            twips = int(gcol.w)
            cols.append(twips / 1440.0)
        return cols
    except Exception:
        return None

def _get_page_usable_width_inches(doc_or_body) -> float:
    try:
        section = doc_or_body.sections[0]
        page_width_in = section.page_width.inches
        left_in = section.left_margin.inches
        right_in = section.right_margin.inches
        return max(0.1, page_width_in - left_in - right_in)
    except AttributeError:
        return 6.5

def _estimate_cell_width_inches(cell, tbl: Table):
    grid_cols = _get_table_grid_col_widths_in_inches(tbl)
    if grid_cols:
        for r in tbl.rows:
            if cell in r.cells:
                col_index = r.cells.index(cell)
                return max(0.1, grid_cols[col_index] - 0.05)
    usable = _get_page_usable_width_inches(tbl._parent)
    ncols = len(tbl.rows[0].cells) if tbl.rows and tbl.rows[0].cells else 2
    return max(0.1, (usable / ncols) - 0.05)

def insert_image_into_cell(cell, tbl: Table, image_bytes: bytes):
    if not image_bytes:
        return
    try:
        with Image.open(io.BytesIO(image_bytes)) as img_check:
            img_check.verify()
    except Exception:
        cell.text = "[Invalid image]"
        return

    cell.text = ""
    par = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell_width_in = _estimate_cell_width_inches(cell, tbl)
    run = par.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Inches(cell_width_in))

def find_paragraph_with_text(doc: Document, placeholder: str):
    for p in doc.paragraphs:
        if placeholder in p.text:
            return p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        return p
    return None

def add_table_after_paragraph(doc: Document, paragraph, rows: int, cols: int) -> Table:
    temp_table = doc.add_table(rows=rows, cols=cols)
    tbl_element = temp_table._tbl
    paragraph._p.addnext(tbl_element)
    return Table(tbl_element, paragraph._parent)

def center_all_cells(tbl: Table):
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def set_equal_column_widths(tbl: Table, total_width_in: float):
    if len(tbl.columns) != 2:
        return
    col_w = total_width_in / 2.0
    for col in tbl.columns:
        for cell in col.cells:
            cell.width = Inches(col_w)

def build_dokumentasi_table_at_placeholder(doc: Document, placeholder: str, items):
    p = find_paragraph_with_text(doc, placeholder)
    if not p:
        p = doc.add_paragraph("")
    _replace_in_paragraph(p, placeholder, "")

    cleaned = []
    for it in items:
        if it.get("image_bytes") or (it.get("caption", "").strip()):
            cleaned.append(it)
    items = cleaned if cleaned else items

    n_items = len(items)
    grid_rows = math.ceil(n_items / 2) if n_items > 0 else 1
    doc_rows = grid_rows * 2
    tbl = add_table_after_paragraph(doc, p, rows=doc_rows, cols=2)
    tbl.autofit = True
    usable = _get_page_usable_width_inches(doc)
    set_equal_column_widths(tbl, usable)
    center_all_cells(tbl)

    idx = 0
    for r in range(grid_rows):
        image_row = tbl.rows[r * 2]
        caption_row = tbl.rows[r * 2 + 1]
        for c in range(2):
            if idx < n_items:
                item = items[idx]
                if item.get("image_bytes"):
                    insert_image_into_cell(image_row.cells[c], tbl, item["image_bytes"])
                else:
                    image_row.cells[c].text = ""
                cap_text = (item.get("caption") or "").strip()
                caption_row.cells[c].text = cap_text
                for pcap in caption_row.cells[c].paragraphs:
                    pcap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                image_row.cells[c].text = ""
                caption_row.cells[c].text = ""
            idx += 1

# -----------------------------------------
# Fungsi pengirim email
# -----------------------------------------

def send_email_with_attachment(
    from_email, password, to_email, smtp_server, smtp_port, subject, body, attachments
):
    """Mengirim email dengan banyak lampiran."""
    try:
        msg = MIMEMultipart()
        msg["From"] = from_email
        msg["To"] = to_email
        msg["Subject"] = subject

        for attachment_bytes, attachment_filename in attachments:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment_bytes)
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f"attachment; filename= {attachment_filename}",
            )
            msg.attach(part)

        server = smtplib.SMTP_SSL(smtp_server, int(smtp_port))
        server.login(from_email, password)
        text = msg.as_string()
        server.sendmail(from_email, to_email, text)
        server.quit()
        return True
    except Exception as e:
        print(f"Error sending email: {e}")
        st.error(f"Gagal mengirim email: {e}")
        return False

# -----------------------------------------
# Streamlit App UI
# -----------------------------------------

st.set_page_config(page_title="Inspection Report Filler", layout="wide")
st.title("Inspection Report Filler")

# Konfigurasi URL template .docx
st.markdown("Aplikasi ini akan secara otomatis mengambil template dari GitHub.")
# Ganti URL ini dengan URL file .docx mentah (raw) Anda di GitHub.
# Anda bisa mendapatkannya dengan membuka file di GitHub, lalu klik tombol "Raw".
template_url = "https://github.com/dpagls/laporan-inspeksi/raw/main/Inspection%20Report%20Template.docx"

# Mengambil template dari URL
try:
    with st.spinner('Mengambil template dari GitHub...'):
        response = requests.get(template_url)
        response.raise_for_status()  # Raises an HTTPError if the status is bad
    template_file = io.BytesIO(response.content)
except requests.exceptions.RequestException as e:
    st.error(f"Gagal mengambil template dari GitHub: {e}. Pastikan URL-nya benar.")
    st.stop()


vessel_list = [
    "MV NAZIHA", "MV AMMAR", "MV NAMEERA", "MV ARFIANIE AYU", "MV SAMI",
    "MV KAREEM", "MV NADHIF", "MV KAYSAN", "MV ABDUL HAMID", "MV NASHALINA",
    "MV GUNALEILA", "MV NUR AWLIYA", "MV NATASCHA", "MV SARAH S", "MV MARIA NASHWAH",
    "MV ZALEHA FITRAT", "MV HAMMADA", "MV KAMADIYA", "MV MUBASYIR", "MV MUHASYIR",
    "MV MUNQIDZ", "MV MUHARRIK", "MV MUMTAZ", "MV UNITAMA LILY", "MT SIL EXPRESS",
    "MT KENCANA EXPRESS", "MV RIMBA EMPAT", "MV MUADZ", "MV MUNIF", "MV RAFA",
    "MT BIO EXPRESS", "MV NOUR MUSHTOFA", "MV FEIZA", "MV MURSYID", "MV AFKAR",
    "MT SELUMA EXPRESS", "MV. AMOLONGO EMRAN", "MV. NIMAOME EMRAN", "MV. SYABIL EMRAN"
]

# Set nilai default untuk rows dokumentasi tanpa menampilkan input
if "dok_rows" not in st.session_state:
    st.session_state.dok_rows = 10

def render_preview_50(file_bytes):
    try:
        with Image.open(io.BytesIO(file_bytes)) as img:
            w, h = img.size
            img_resized = img.resize((max(1, w // 2), max(1, h // 2)))
        st.image(img_resized, use_column_width=False)
    except Exception:
        st.warning("Format gambar tidak valid.")

st.subheader("FOTOHALUAN")
foto_haluan_file = st.file_uploader(
    "Upload FOTOHALUAN image", type=["jpg", "jpeg", "png"], key="foto_haluan"
)

if foto_haluan_file is not None:
    st.session_state["foto_haluan_bytes"] = foto_haluan_file.getvalue()

if "foto_haluan_bytes" in st.session_state:
    render_preview_50(st.session_state["foto_haluan_bytes"])

st.markdown("---")

st.subheader("Vessel Details")
col1, col2 = st.columns(2, gap="large")
with col1:
    vessel_name = st.selectbox("Vessel (*VESSEL*)", options=vessel_list)
    imo = st.text_input("IMO (*IMO*)")
with col2:
    ship_type = st.selectbox("Type (*TYPE*)", options=["Bulk Carrier", "Tanker"])
    callsign = st.text_input("Callsign (*CALLSIGN*)")

place = st.text_input("Place (*PLACEDATE*)", placeholder="e.g., Jakarta")
survey_date = st.date_input("Date (*PLACEDATE*)", value=date.today())
master = st.text_input("Master (*MASTER*)")
surveyor = st.text_input("Surveyor (*SURVEYOR*)")

st.markdown("---")

st.subheader("DOKUMENTASI")
dok_items = []
row_pairs = [(i, i + 1) for i in range(0, st.session_state.dok_rows, 2)]

for left_idx, right_idx in row_pairs:
    col_left, col_right = st.columns(2, gap="large")
    with col_left:
        img_key_left = f"dok_img_{left_idx}_0"
        cap_key_left = f"dok_cap_{left_idx}_0"
        file_left = st.file_uploader(
            f"Row {left_idx + 1} - Left Image",
            type=["jpg", "jpeg", "png"],
            key=img_key_left
        )
        if file_left is not None:
            st.session_state[img_key_left + "_bytes"] = file_left.getvalue()
        if st.session_state.get(img_key_left + "_bytes"):
            render_preview_50(st.session_state[img_key_left + "_bytes"])
        caption_left = st.text_input(f"Row {left_idx + 1} - Caption", key=cap_key_left)
        dok_items.append({
            "image_bytes": st.session_state.get(img_key_left + "_bytes", None),
            "caption": caption_left or ""
        })
    with col_right:
        if right_idx < st.session_state.dok_rows:
            img_key_right = f"dok_img_{right_idx}_1"
            cap_key_right = f"dok_cap_{right_idx}_1"
            file_right = st.file_uploader(
                f"Row {right_idx + 1} - Right Image",
                type=["jpg", "jpeg", "png"],
                key=img_key_right
            )
            if file_right is not None:
                st.session_state[img_key_right + "_bytes"] = file_right.getvalue()
            if st.session_state.get(img_key_right + "_bytes"):
                render_preview_50(st.session_state[img_key_right + "_bytes"])
            caption_right = st.text_input(f"Row {right_idx + 1} - Caption", key=cap_key_right)
            dok_items.append({
                "image_bytes": st.session_state.get(img_key_right + "_bytes", None),
                "caption": caption_right or ""
            })

def add_dok_row():
    st.session_state.dok_rows += 2
    st.rerun()

st.markdown("---")
st.button("➕ Tambah Row Dokumentasi", on_click=add_dok_row)

st.markdown("---")

# -----------------------------------------
# Konfigurasi Email
# -----------------------------------------
st.subheader("Email Penerima")
st.info("Masukkan email penerima di bawah ini. Laporan akan langsung dikirim setelah di-generate.")

# Konfigurasi SMTP diatur di sini, tidak terlihat oleh pengguna
email_sender = "fajar@dpagls.my.id"
email_password = "Rahasia100%"
smtp_server = "mail.dpagls.my.id"
smtp_port = "465"

email_to_send = st.text_input("Email Penerima")

# -----------------------------------------
# Generate Report
# -----------------------------------------
if st.button("📝 Generate Report"):
    if not email_to_send:
        st.error("Silakan masukkan alamat email penerima.")
        st.stop()
    
    try:
        doc = Document(template_file)
    except Exception:
        st.error("Template .docx tidak valid atau rusak.")
        st.stop()

    replace_placeholder_everywhere(doc, "*VESSEL*", vessel_name)
    replace_placeholder_everywhere(doc, "*IMO*", imo)
    replace_placeholder_everywhere(doc, "*TYPE*", ship_type)
    replace_placeholder_everywhere(doc, "*CALLSIGN*", callsign)
    replace_placeholder_everywhere(doc, "*PLACEDATE*", f"{place}, {survey_date.strftime('%d %B %Y')}")
    replace_placeholder_everywhere(doc, "*MASTER*", master)
    replace_placeholder_everywhere(doc, "*SURVEYOR*", surveyor)

    foto_bytes = st.session_state.get("foto_haluan_bytes")
    if foto_bytes:
        cell, tbl, _, _ = find_cell_with_text(doc, "*FOTOHALUAN*")
        if cell:
            insert_image_into_cell(cell, tbl, foto_bytes)
            replace_placeholder_everywhere(doc, "*FOTOHALUAN*", "")
        else:
            replace_placeholder_everywhere(doc, "*FOTOHALUAN*", "")

    build_dokumentasi_table_at_placeholder(doc, "*DOKUMENTASI*", dok_items)

    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    # Nama file
    base_filename = f"{survey_date.strftime('%Y.%m.%d')} {vessel_name} Inspection Report"
    docx_filename = f"{base_filename}.docx"

    # Kirim email
    st.write("Mengirim laporan melalui email...")
    attachments_list = [
        (docx_buffer.getvalue(), docx_filename)
    ]
    success = send_email_with_attachment(
        email_sender,
        email_password,
        email_to_send,
        smtp_server,
        smtp_port,
        f"Laporan Inspeksi: {vessel_name}",
        f"Terlampir laporan inspeksi kapal {vessel_name} dalam format DOCX.",
        attachments_list
    )
    if success:
        st.success(f"Laporan berhasil dikirim ke {email_to_send}!")
    
    # Tombol Download
    st.download_button(
        label="📄 Download Laporan (DOCX)",
        data=docx_buffer,
        file_name=docx_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
